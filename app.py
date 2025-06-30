import os
import uvicorn
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse
from docx import Document
import re
from io import BytesIO
import gc
app = FastAPI()

def has_color(run):
    return run.font.color.rgb is not None

def has_highlight_color(run):
    return run.font.highlight_color is not None

def has_bold(run):
    return run.font.bold is not None

def process_quiz_format(doc):
    """Process document with quiz format (1 question, 4 answers)"""
    data = []
    
    paragraphs = iter(doc.paragraphs)
    for paragraph in paragraphs:
        if paragraph.text.startswith("Câu") or paragraph.text.endswith("?") or paragraph.text.endswith(":"):
            # Remove numbering pattern like "1. ", "2. ", etc. at the beginning of the question
            question_text = paragraph.text.strip()
            question_text = re.sub(r"^\d+\.\s*", "", question_text)
            # Also remove "Câu XX: " pattern if present
            question_text = re.sub(r"^Câu\s+\d+\s*[:.\s]\s*", "", question_text)
            
            question_data = {
                "question": question_text,
                "answers": [],
                "correct": ""
            }

            options = []
            correct_answer = ""

            for i in range(4):  # Assuming 4 options for each question
                try:
                    option_paragraph = next(paragraphs)
                    if not option_paragraph.text.strip():  # Skip empty paragraphs
                        continue
                        
                    # Clean up option text (remove A., B., etc. prefixes)
                    option_text = option_paragraph.text.strip()
                    option_text = re.sub(r"^[A-D]\.\s*", "", option_text)
                    
                    options.append(option_text)
                    
                    # Check if this option is marked as correct (bold or highlighted)
                    is_correct = False
                    for option_run in option_paragraph.runs:
                        if has_bold(option_run) or has_highlight_color(option_run) or has_color(option_run):
                            correct_answer = i  # Use index position (0, 1, 2, 3)
                            is_correct = True
                            break
                            
                except StopIteration:
                    break

            question_data["answers"] = options
            question_data["correct"] = correct_answer
            data.append(question_data)
    
    return data

def process_qa_format(doc):
    """Process document with Q&A format (1 question, 1 answer)"""
    data = []
    
    current_question = None
    current_answer = None
    i = 0
    paragraphs = list(doc.paragraphs)

    while i < len(paragraphs):
        text = paragraphs[i].text.strip()
        if not text:  # Skip empty paragraphs
            i += 1
            continue
            
        # Check if this is a question (starts with number or ends with ? or :)
        if text.endswith("?") or text.endswith(":") or re.match(r"^\d+\.", text):
            # Save previous question if exists
            if current_question and current_answer:
                data.append({
                    "question": current_question,
                    "answer": current_answer
                })
                
            # Start new question - remove numbering
            current_question = text
            # Remove numbering pattern like "1. ", "2. ", etc.
            current_question = re.sub(r"^\d+\.\s*", "", current_question)
            # Also remove "Câu XX: " pattern if present
            current_question = re.sub(r"^Câu\s+\d+\s*[:.\s]\s*", "", current_question)
            
            # Look ahead for the answer
            current_answer = ""
            j = i + 1
            while j < len(paragraphs):
                answer_text = paragraphs[j].text.strip()
                if not answer_text:  # Skip empty paragraphs
                    j += 1
                    continue
                
                if answer_text.endswith("?") or answer_text.endswith(":") or re.match(r"^\d+\.", answer_text):
                    # This is the next question, stop here
                    break
                
                # This is the answer
                current_answer = answer_text
                # Check if text has formatting that indicates it's a correct answer
                for run in paragraphs[j].runs:
                    if has_bold(run) or has_highlight_color(run) or has_color(run):
                        current_answer = answer_text
                
                # Move the main index to skip the answer we just processed
                i = j
                break
                
            if not current_answer and j < len(paragraphs):
                # If we didn't find a suitable answer, use the next non-empty paragraph
                current_answer = "No answer provided"
        
        i += 1
    
    # Add the last question
    if current_question and current_answer:
        data.append({
            "question": current_question,
            "answer": current_answer,
        })
    
    return data

@app.get("/")
async def read_root():
    return {
        "message": "DOCX to JSON Converter API",
        "endpoints": {
            "/quiz": "POST - Convert DOCX to quiz format (1 question, 4 answers)",
            "/api/qa": "POST - Convert DOCX to Q&A format (1 question, 1 answer)"
        },
        "usage": "Upload .docx file using multipart/form-data"
    }

@app.post('/quiz')
async def convert_quiz(file: UploadFile = File(...)):
    """API endpoint for quiz format (1 question, 4 answers with correct answer marked)"""
    try:
        # Validate file type
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="File must be a .docx file")
        
        # Read file content into memory
        file_content = await file.read()
        
        # Process document from memory
        doc = Document(BytesIO(file_content))
        result = process_quiz_format(doc)
        
        # Clear memory
        del file_content
        del doc
        gc.collect()
        
        return JSONResponse(content=result)
        
    except Exception as e:
        # Clear memory in case of error
        gc.collect()
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@app.post('/api/qa')
async def convert_qa(file: UploadFile = File(...)):
    """API endpoint for Q&A format (1 question, 1 answer)"""
    try:
        # Validate file type
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="File must be a .docx file")
        
        # Read file content into memory
        file_content = await file.read()
        
        # Process document from memory
        doc = Document(BytesIO(file_content))
        result = process_qa_format(doc)
        
        # Clear memory
        del file_content
        del doc
        gc.collect()
        
        return JSONResponse(content=result)
        
    except Exception as e:
        # Clear memory in case of error
        gc.collect()
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000)) # Mặc định là 8000 cho FastAPI
    uvicorn.run(app, host="0.0.0.0", port=port)