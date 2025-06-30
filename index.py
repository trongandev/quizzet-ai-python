from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
import re
import os
import uuid
import threading
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from queue import Queue
import logging
from io import BytesIO

app = Flask(__name__)
CORS(app, origins="*", methods=["GET", "POST", "OPTIONS"], allow_headers=["Content-Type", "Authorization"])

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Thread pool for handling concurrent requests
MAX_WORKERS = 10  # Adjust based on your server capacity
executor = ThreadPoolExecutor(max_workers=MAX_WORKERS)

# Queue for managing processing requests
processing_queue = Queue(maxsize=50)  # Max 50 pending requests
active_tasks = {}  # Track active processing tasks
task_results = {}  # Store completed task results
task_lock = threading.Lock()

class TaskStatus:
    PENDING = "pending"
    PROCESSING = "processing" 
    COMPLETED = "completed"
    FAILED = "failed"

def has_color(run):
    return run.font.color.rgb is not None

def has_highlight_color(run):
    return run.font.highlight_color is not None

def has_bold(run):
    return run.font.bold is not None

def process_quiz_format(doc):
    """Process document with quiz format (1 question, 4 answers)"""
    data = []
    
    paragraphs = iter(doc.paragraphs)
    for paragraph in paragraphs:
        if paragraph.text.startswith("Câu") or paragraph.text.endswith("?") or paragraph.text.endswith(":"):
            # Remove numbering pattern like "1. ", "2. ", etc. at the beginning of the question
            question_text = paragraph.text.strip()
            question_text = re.sub(r"^\d+\.\s*", "", question_text)
            # Also remove "Câu XX: " pattern if present
            question_text = re.sub(r"^Câu\s+\d+\s*[:.\s]\s*", "", question_text)
            
            question_data = {
                "question": question_text,
                "answers": [],
                "correct": ""
            }

            options = []
            correct_answer = ""

            for i in range(4):  # Assuming 4 options for each question
                try:
                    option_paragraph = next(paragraphs)
                    if not option_paragraph.text.strip():  # Skip empty paragraphs
                        continue
                        
                    # Clean up option text (remove A., B., etc. prefixes)
                    option_text = option_paragraph.text.strip()
                    option_text = re.sub(r"^[A-D]\.\s*", "", option_text)
                    
                    options.append(option_text)
                    
                    # Check if this option is marked as correct (bold or highlighted)
                    is_correct = False
                    for option_run in option_paragraph.runs:
                        if has_bold(option_run) or has_highlight_color(option_run) or has_color(option_run):
                            correct_answer = i  # Use index position (0, 1, 2, 3)
                            is_correct = True
                            break
                            
                except StopIteration:
                    break

            question_data["answers"] = options
            question_data["correct"] = correct_answer
            data.append(question_data)
    
    return data

def process_qa_format(doc):
    """Process document with Q&A format (1 question, 1 answer)"""
    data = []
    
    current_question = None
    current_answer = None
    i = 0
    paragraphs = list(doc.paragraphs)

    while i < len(paragraphs):
        text = paragraphs[i].text.strip()
        if not text:  # Skip empty paragraphs
            i += 1
            continue
            
        # Check if this is a question (starts with number or ends with ? or :)
        if text.endswith("?") or text.endswith(":") or re.match(r"^\d+\.", text):
            # Save previous question if exists
            if current_question and current_answer:
                data.append({
                    "question": current_question,
                    "answer": current_answer
                })
                
            # Start new question - remove numbering
            current_question = text
            # Remove numbering pattern like "1. ", "2. ", etc.
            current_question = re.sub(r"^\d+\.\s*", "", current_question)
            # Also remove "Câu XX: " pattern if present
            current_question = re.sub(r"^Câu\s+\d+\s*[:.\s]\s*", "", current_question)
            
            # Look ahead for the answer
            current_answer = ""
            j = i + 1
            while j < len(paragraphs):
                answer_text = paragraphs[j].text.strip()
                if not answer_text:  # Skip empty paragraphs
                    j += 1
                    continue
                
                if answer_text.endswith("?") or answer_text.endswith(":") or re.match(r"^\d+\.", answer_text):
                    # This is the next question, stop here
                    break
                
                # This is the answer
                current_answer = answer_text
                # Check if text has formatting that indicates it's a correct answer
                for run in paragraphs[j].runs:
                    if has_bold(run) or has_highlight_color(run) or has_color(run):
                        current_answer = answer_text
                
                # Move the main index to skip the answer we just processed
                i = j
                break
                
            if not current_answer and j < len(paragraphs):
                # If we didn't find a suitable answer, use the next non-empty paragraph
                current_answer = "No answer provided"
        
        i += 1
    
    # Add the last question
    if current_question and current_answer:
        data.append({
            "question": current_question,
            "answer": current_answer,
        })
    
    return data

def process_document_async(task_id, file_data, process_type):
    """Process document asynchronously"""
    try:
        logger.info(f"Starting processing task {task_id}")
        
        # Update task status
        with task_lock:
            active_tasks[task_id] = {
                "status": TaskStatus.PROCESSING,
                "start_time": time.time(),
                "progress": 0
            }
        
        # Process document from memory
        doc = Document(BytesIO(file_data))
        
        if process_type == "quiz":
            result = process_quiz_format(doc)
        else:  # qa
            result = process_qa_format(doc)
        
        # Store result
        with task_lock:
            task_results[task_id] = {
                "status": TaskStatus.COMPLETED,
                "result": result,
                "completion_time": time.time(),
                "processing_time": time.time() - active_tasks[task_id]["start_time"]
            }
            # Remove from active tasks
            del active_tasks[task_id]
        
        logger.info(f"Completed processing task {task_id}")
        
    except Exception as e:
        logger.error(f"Error processing task {task_id}: {str(e)}")
        with task_lock:
            task_results[task_id] = {
                "status": TaskStatus.FAILED,
                "error": str(e),
                "completion_time": time.time()
            }
            if task_id in active_tasks:
                del active_tasks[task_id]

def cleanup_old_results():
    """Clean up old results to prevent memory leaks"""
    current_time = time.time()
    with task_lock:
        # Remove results older than 1 hour
        expired_tasks = [
            task_id for task_id, result in task_results.items()
            if current_time - result.get("completion_time", 0) > 3600
        ]
        for task_id in expired_tasks:
            del task_results[task_id]
            logger.info(f"Cleaned up expired task result: {task_id}")



# Background thread for cleanup
def cleanup_worker():
    while True:
        time.sleep(600)  # Run every 10 minutes
        cleanup_old_results()

cleanup_thread = threading.Thread(target=cleanup_worker, daemon=True)
cleanup_thread.start()

@app.errorhandler(405)
def method_not_allowed(error):
    return jsonify({
        "error": "Method Not Allowed",
        "message": "The method is not allowed for the requested URL",
        "allowed_methods": ["GET", "POST"],
        "url": request.url,
        "method": request.method
    }), 405

@app.errorhandler(404)
def not_found(error):
    return jsonify({
        "error": "Not Found",
        "message": "The requested URL was not found",
        "url": request.url
    }), 404

@app.route('/test', methods=['GET', 'POST'])
def test_endpoint():
    """Simple test endpoint"""
    return jsonify({
        "status": "success",
        "message": "API is working correctly",
        "method": request.method,
        "timestamp": time.time()
    })

@app.route('/api/quiz', methods=['POST', 'GET'])
def convert_quiz():
    """API endpoint for quiz format (1 question, 4 answers with correct answer marked)"""
    if request.method == 'GET':
        return jsonify({
            "message": "Quiz API endpoint is working",
            "method": "POST",
            "description": "Upload a .docx file to convert to quiz format"
        })
    return process_document_endpoint("quiz")

@app.route('/api/qa', methods=['POST', 'GET'])
def convert_qa():
    """API endpoint for Q&A format (1 question, 1 answer)"""
    if request.method == 'GET':
        return jsonify({
            "message": "Q&A API endpoint is working", 
            "method": "POST",
            "description": "Upload a .docx file to convert to Q&A format"
        })
    return process_document_endpoint("qa")

def process_document_endpoint(process_type):
    """Common endpoint logic for processing documents"""
    try:
        # Check if file is present in request
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        
        # Check if file has .docx extension
        if not file.filename.endswith('.docx'):
            return jsonify({"error": "File must be a .docx file"}), 400
        
        # Check if we can handle more requests
        if len(active_tasks) >= MAX_WORKERS:
            return jsonify({
                "error": "Server is busy. Please try again later.",
                "active_tasks": len(active_tasks),
                "max_capacity": MAX_WORKERS
            }), 503
        
        # Generate unique task ID
        task_id = str(uuid.uuid4())
        
        # Read file data into memory
        file_data = file.read()
        
        # Check if client wants synchronous or asynchronous processing
        async_mode = request.args.get('async', 'false').lower() == 'true'
        
        if async_mode:
            # Submit to thread pool for async processing
            executor.submit(process_document_async, task_id, file_data, process_type)
            
            return jsonify({
                "task_id": task_id,
                "status": TaskStatus.PENDING,
                "message": "Document submitted for processing",
                "check_status_url": f"/api/status/{task_id}"
            }), 202
        else:
            # Process synchronously (original behavior)
            try:
                doc = Document(BytesIO(file_data))
                
                if process_type == "quiz":
                    result = process_quiz_format(doc)
                else:  # qa
                    result = process_qa_format(doc)
                
                return jsonify(result)
                
            except Exception as e:
                raise e
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/status/<task_id>', methods=['GET'])
def get_task_status(task_id):
    """Get status of an async task"""
    with task_lock:
        # Check if task is still processing
        if task_id in active_tasks:
            return jsonify({
                "task_id": task_id,
                "status": active_tasks[task_id]["status"],
                "processing_time": time.time() - active_tasks[task_id]["start_time"]
            })
        
        # Check if task is completed
        if task_id in task_results:
            result = task_results[task_id].copy()
            result["task_id"] = task_id
            return jsonify(result)
        
        # Task not found
        return jsonify({
            "error": "Task not found",
            "task_id": task_id
        }), 404

@app.route('/api/server-status', methods=['GET'])
def get_server_status():
    """Get server status and load information"""
    with task_lock:
        return jsonify({
            "active_tasks": len(active_tasks),
            "max_workers": MAX_WORKERS,
            "completed_tasks": len(task_results),
            "server_load": f"{len(active_tasks)}/{MAX_WORKERS}",
            "queue_status": "healthy" if len(active_tasks) < MAX_WORKERS else "busy"
        })

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({"status": "healthy", "message": "DOCX to JSON API is running"}), 200

@app.route('/', methods=['GET'])
def home():
    """Home endpoint with API documentation"""
    return jsonify({
        "message": "DOCX to JSON Converter API with Memory Processing",
        "endpoints": {
            "/api/quiz": {
                "method": "POST",
                "description": "Convert DOCX to JSON with quiz format (1 question, 4 answers)",
                "parameters": "file: .docx file, async: true/false (optional)"
            },
            "/api/qa": {
                "method": "POST", 
                "description": "Convert DOCX to JSON with Q&A format (1 question, 1 answer)",
                "parameters": "file: .docx file, async: true/false (optional)"
            },
            "/api/status/<task_id>": {
                "method": "GET",
                "description": "Get status of async task"
            },
            "/api/server-status": {
                "method": "GET",
                "description": "Get server load and status information"
            },
            "/api/health": {
                "method": "GET",
                "description": "Health check endpoint"
            }
        },
        "usage": {
            "sync_request": "POST /api/quiz with file",
            "async_request": "POST /api/quiz?async=true with file",
            "check_status": "GET /api/status/{task_id}"
        },
        "features": {
            "memory_processing": "Files are processed in memory without saving to disk",
            "railway_compatible": "Optimized for Railway deployment",
            "async_support": "Supports both sync and async processing"
        }
    })

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5999))
    app.run(debug=False, host='0.0.0.0', port=port, threaded=True)
